#!/usr/bin/env python3
"""Render constrained-markdown source docs into FEB-styled .docx files.

Supported markdown subset (everything the SN6 Resources sources use):
  # / ## / ### headings, pipe tables, - bullets, 1. numbered lists,
  **bold** / *italic* / `code` inline, > blockquote (rendered as CAUTION),
  --- horizontal rule (ignored), plain paragraphs.

Usage:
  build_docx.py SRC.md OUT.docx [--title-override "..."]
  build_docx.py --all           # build every known source in the repo layout
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

FEB_BLUE = RGBColor(0x00, 0x32, 0x62)   # Berkeley blue
FEB_GOLD = RGBColor(0xFD, 0xB5, 0x15)   # California gold
BODY_FONT = "Calibri"

ROOT = Path(__file__).resolve().parent.parent


def _style_base(doc):
    st = doc.styles["Normal"]
    st.font.name = BODY_FONT
    st.font.size = Pt(10.5)
    for lvl, size in ((1, 17), (2, 13.5), (3, 11.5)):
        h = doc.styles[f"Heading {lvl}"]
        h.font.name = BODY_FONT
        h.font.size = Pt(size)
        h.font.color.rgb = FEB_BLUE
        h.font.bold = True


def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def _add_runs(par, text):
    for chunk in INLINE_RE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            r = par.add_run(chunk[2:-2])
            r.bold = True
        elif chunk.startswith("*") and chunk.endswith("*") and len(chunk) > 2:
            r = par.add_run(chunk[1:-1])
            r.italic = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            r = par.add_run(chunk[1:-1])
            r.font.name = "Courier New"
        else:
            par.add_run(chunk)


def _clean(text):
    # strip markdown links [text](url) -> text (url)
    return re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)


def _add_table(doc, rows):
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
            continue  # separator row
        parsed.append(cells)
    if not parsed:
        return
    ncols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(parsed):
        for j in range(ncols):
            cell = table.cell(i, j)
            text = _clean(row[j]) if j < len(row) else ""
            par = cell.paragraphs[0]
            _add_runs(par, text)
            for r in par.runs:
                r.font.size = Pt(9.5)
            if i == 0:
                _shade(cell, "003262")
                for r in par.runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.bold = True
    doc.add_paragraph()


PHOTO_RE = re.compile(r"\[PHOTO:([^\]]+)\]")


def _add_photo_placeholder(doc, caption):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _shade(cell, "F2F2F2")
    par = cell.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run(f"\n[ PHOTO: {caption.strip()} ]\n\n")
    r.italic = True
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_paragraph()


def build(src: Path, out: Path):
    lines = src.read_text().splitlines()
    doc = Document()
    _style_base(doc)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.8)
        section.left_margin = section.right_margin = Inches(0.9)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped or stripped == "---":
            i += 1
            continue
        if stripped.startswith("|"):
            j = i
            while j < len(lines) and lines[j].strip().startswith("|"):
                j += 1
            _add_table(doc, lines[i:j])
            i = j
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            lvl = min(len(m.group(1)), 3)
            doc.add_heading("", level=lvl)
            _add_runs(doc.paragraphs[-1], _clean(m.group(2)))
            i += 1
            continue
        pm = PHOTO_RE.fullmatch(stripped)
        if pm:
            _add_photo_placeholder(doc, pm.group(1))
            i += 1
            continue
        if stripped.startswith(">"):
            par = doc.add_paragraph()
            r = par.add_run("⚠ CAUTION: ")
            r.bold = True
            _add_runs(par, _clean(stripped.lstrip("> ")))
            par.paragraph_format.left_indent = Inches(0.3)
            i += 1
            continue
        if re.match(r"^[-*]\s+", stripped):
            par = doc.add_paragraph(style="List Bullet")
            _add_runs(par, _clean(re.sub(r"^[-*]\s+", "", stripped)))
            i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            par = doc.add_paragraph(style="List Number")
            _add_runs(par, _clean(re.sub(r"^\d+\.\s+", "", stripped)))
            i += 1
            continue
        par = doc.add_paragraph()
        _add_runs(par, _clean(stripped))
        i += 1

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"built {out.relative_to(ROOT) if out.is_relative_to(ROOT) else out}")


KNOWN = [
    ("01 Pain Points and Improvements/src/pain-points.md",
     "01 Pain Points and Improvements/SN5 Pain Points and Improvements.docx"),
]


def build_all():
    # pain points
    for src, out in KNOWN:
        if (ROOT / src).exists():
            build(ROOT / src, ROOT / out)
    # CS standards: src/CS-*.md -> "<id> <title>.docx" from first heading
    cs_src = ROOT / "02 CS Standards" / "src"
    for md in sorted(cs_src.glob("CS-*.md")):
        first = md.read_text().splitlines()[0].lstrip("# ").strip()
        name = re.sub(r"[/:]", "-", first)
        out = ROOT / "02 CS Standards" / f"{name}.docx"
        build(md, out)
    tmpl = ROOT / "02 CS Standards" / "template" / "CS-Template.md"
    if tmpl.exists():
        build(tmpl, tmpl.with_suffix(".docx"))


if __name__ == "__main__":
    if "--all" in sys.argv:
        build_all()
    elif len(sys.argv) >= 3:
        build(Path(sys.argv[1]), Path(sys.argv[2]))
    else:
        sys.exit(__doc__)
